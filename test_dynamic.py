import requests
import docx
import os

# Create Case 1: Civil Engineering
doc_cv = docx.Document()
doc_cv.add_paragraph('Experienced Civil Engineer skilled in AutoCAD, STAAD, and Structural Analysis.')
case1_path = os.path.abspath('case1_civil.docx')
doc_cv.save(case1_path)

# Create Case 2: ML / Python
doc_ml = docx.Document()
doc_ml.add_paragraph('Machine Learning Engineer proficient in Python, PyTorch, Keras, and deep learning LLMs.')
case2_path = os.path.abspath('case2_ml.docx')
doc_ml.save(case2_path)

def test_endpoint(file_path, jd=""):
    print(f"--- Testing {file_path} ---")
    try:
        with open(file_path, 'rb') as f:
            files = {'resume': f}
            data = {'job_description': jd}
            response = requests.post('http://127.0.0.1:5000/analyze', files=files, data=data)
            html = response.text
            # Extract Role
            import re
            role_match = re.search(r'Predicted Dominant Role</p>\s*<h3 class="display-6 fw-bold text-white mb-4">([^<]+)</h3>', html)
            if role_match:
                print("Predicted Role:", role_match.group(1))
            else:
                print("Role not found in HTML")
            
            # Extract ATS Score
            ats_match = re.search(r'id="atsScoreChart" data-score="([^"]+)"', html)
            if ats_match:
                print("ATS Match Score:", ats_match.group(1))
    except Exception as e:
        print("Error:", e)
        
test_endpoint(case1_path)
test_endpoint(case2_path)
test_endpoint(case2_path, jd="Requires AWS and Docker infrastructure")

os.remove(case1_path)
os.remove(case2_path)

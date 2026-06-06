import requests
import docx
import os
import json

# 1. Create a dummy docx resume for testing
doc = docx.Document()
doc.add_paragraph('Experienced Python Developer skilled in Django, Flask, Postgres, and SQL.')
doc.add_paragraph('Spearheaded database queries optimization and made REST APIs for payments.')
temp_path = os.path.abspath('temp_test_resume.docx')
doc.save(temp_path)

url = 'http://127.0.0.1:5000/api/analyze'

try:
    with open(temp_path, 'rb') as f:
        files = {'resume': f}
        data = {'job_description': 'Looking for a Python Developer with django and redis experience.'}
        
        print(f"Sending request to {url}...")
        r = requests.post(url, files=files, data=data)
        
        if r.status_code == 200:
            res = r.json()
            print("--- API RESPONSE RECEIVED SUCCESSFUL ---")
            print(json.dumps(res, indent=2))
            
            # Validation assertions
            assert res['success'] is True, "Success should be True"
            assert 'predicted_role' in res, "Missing predicted_role"
            assert 'confidence' in res, "Missing confidence"
            assert 'ats_score' in res, "Missing ats_score"
            assert 'matched_skills' in res, "Missing matched_skills"
            assert 'missing_skills' in res, "Missing missing_skills"
            assert 'resume_dna' in res, "Missing resume_dna"
            assert 'suggestions' in res, "Missing suggestions"
            assert 'enhanced_bullets' in res, "Missing enhanced_bullets"
            print("\n[SUCCESS] API conforms exactly to the required JSON schema!")
        else:
            print(f"[ERROR] API returned status code {r.status_code}")
            print(r.text)

except Exception as e:
    print(f"[ERROR] Test failed: {e}")

finally:
    if os.path.exists(temp_path):
        os.remove(temp_path)

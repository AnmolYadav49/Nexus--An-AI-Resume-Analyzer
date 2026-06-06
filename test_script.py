import docx
import os

doc = docx.Document()
doc.add_heading('Anmol Yadav - Python Developer', 0)
doc.add_paragraph('Experienced Python Developer with 4 years of background in creating scalable applications using Django, Flask, and FastAPI.')
doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Django, Flask, REST APIs, Docker, Linux, SQL, Postgres.')
doc.add_heading('Projects', level=1)
doc.add_paragraph('Built a scalable microservice architecture that improved response times by 30%.')

file_path = os.path.join(os.path.dirname(__file__), 'test_resume.docx')
doc.save(file_path)
print(f"Created dummy resume at {file_path}")
